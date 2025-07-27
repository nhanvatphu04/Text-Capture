from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

class TextProcessor:
    def __init__(self, text: str):
        self.text = text
        self.document = Document()
    
    def get_text(self) -> str:
        return self.text
    
    def capitalize_sentences(self) -> str:
        """
        In hoa chữ cái đầu của mỗi câu
        """
        if not self.text:
            return self.text
        
        # Tách văn bản thành các câu
        sentences = self.text.split('. ')
        
        # Xử lý từng câu
        processed_sentences = []
        for sentence in sentences:
            if sentence.strip():
                # In hoa chữ cái đầu và giữ nguyên phần còn lại
                processed_sentence = sentence.strip().capitalize()
                processed_sentences.append(processed_sentence)
        
        # Nối lại các câu với dấu chấm
        result = '. '.join(processed_sentences)
        
        # Đảm bảo kết thúc bằng dấu chấm nếu văn bản gốc có
        if self.text.strip().endswith('.'):
            result += '.'
        
        return result
    
    def to_lowercase(self) -> str:
        """
        Chuyển toàn bộ văn bản thành chữ thường
        """
        return self.text.lower()
    
    def to_uppercase(self) -> str:
        """
        Chuyển toàn bộ văn bản thành chữ hoa
        """
        return self.text.upper()
    
    def add_underline(self) -> str:
        """
        Thêm gạch chân cho toàn bộ văn bản
        """
        if not self.text:
            return self.text
        
        # Tạo paragraph với gạch chân
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(self.text)
        run.underline = True
        
        return self.text  # Trả về text gốc, định dạng được lưu trong document
    
    def add_strikethrough(self) -> str:
        """
        Thêm gạch ngang cho toàn bộ văn bản
        """
        if not self.text:
            return self.text
        
        # Tạo paragraph với gạch ngang
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(self.text)
        run.font.strike = True
        
        return self.text  # Trả về text gốc, định dạng được lưu trong document
    
    def add_underline_to_words(self) -> str:
        """
        Thêm gạch chân cho từng từ riêng biệt
        """
        if not self.text:
            return self.text
        
        # Tách văn bản thành các từ
        words = self.text.split()
        
        # Tạo paragraph với gạch chân cho từng từ
        paragraph = self.document.add_paragraph()
        for i, word in enumerate(words):
            run = paragraph.add_run(word)
            run.underline = True
            if i < len(words) - 1:
                paragraph.add_run(" ")  # Thêm khoảng trắng
        
        return self.text  # Trả về text gốc, định dạng được lưu trong document
    
    def add_strikethrough_to_words(self) -> str:
        """
        Thêm gạch ngang cho từng từ riêng biệt
        """
        if not self.text:
            return self.text
        
        # Tách văn bản thành các từ
        words = self.text.split()
        
        # Tạo paragraph với gạch ngang cho từng từ
        paragraph = self.document.add_paragraph()
        for i, word in enumerate(words):
            run = paragraph.add_run(word)
            run.font.strike = True
            if i < len(words) - 1:
                paragraph.add_run(" ")  # Thêm khoảng trắng
        
        return self.text  # Trả về text gốc, định dạng được lưu trong document
    
    def convert_underline_to_strikethrough(self) -> str:
        """
        Chuyển đổi từ gạch chân sang gạch ngang
        """
        if not self.text:
            return self.text
        
        # Tạo paragraph với gạch ngang
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(self.text)
        run.font.strike = True
        
        return self.text  # Trả về text gốc, định dạng được lưu trong document
    
    def convert_strikethrough_to_underline(self) -> str:
        """
        Chuyển đổi từ gạch ngang sang gạch chân
        """
        if not self.text:
            return self.text
        
        # Tạo paragraph với gạch chân
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(self.text)
        run.underline = True
        
        return self.text  # Trả về text gốc, định dạng được lưu trong document
    
    def remove_underline(self) -> str:
        """
        Loại bỏ tất cả gạch chân khỏi văn bản
        """
        if not self.text:
            return self.text
        
        # Tạo paragraph không có định dạng
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(self.text)
        run.underline = False
        
        return self.text  # Trả về text gốc, định dạng được lưu trong document
    
    def remove_strikethrough(self) -> str:
        """
        Loại bỏ tất cả gạch ngang khỏi văn bản
        """
        if not self.text:
            return self.text
        
        # Tạo paragraph không có định dạng
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(self.text)
        run.font.strike = False
        
        return self.text  # Trả về text gốc, định dạng được lưu trong document
    
    def remove_all_formatting(self) -> str:
        """
        Loại bỏ tất cả định dạng gạch chân và gạch ngang
        """
        if not self.text:
            return self.text
        
        # Tạo paragraph không có định dạng
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(self.text)
        run.underline = False
        run.font.strike = False
        
        return self.text  # Trả về text gốc, định dạng được lưu trong document
    
    def toggle_underline(self) -> str:
        """
        Bật/tắt gạch chân cho văn bản
        """
        if not self.text:
            return self.text
        
        # Tạo paragraph với gạch chân
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(self.text)
        run.underline = True
        
        return self.text  # Trả về text gốc, định dạng được lưu trong document
    
    def toggle_strikethrough(self) -> str:
        """
        Bật/tắt gạch ngang cho văn bản
        """
        if not self.text:
            return self.text
        
        # Tạo paragraph với gạch ngang
        paragraph = self.document.add_paragraph()
        run = paragraph.add_run(self.text)
        run.font.strike = True
        
        return self.text  # Trả về text gốc, định dạng được lưu trong document
    
    def save_document(self, filename: str = "formatted_text.docx") -> str:
        """
        Lưu document với định dạng vào file
        """
        try:
            self.document.save(filename)
            return f"Document saved as {filename}"
        except Exception as e:
            return f"Error saving document: {e}"
    
    def clear_document(self):
        """
        Xóa tất cả nội dung trong document
        """
        self.document = Document()
    
    def set_text(self, text: str) -> None:
        """
        Cập nhật văn bản mới
        """
        self.text = text
        self.clear_document()  # Xóa document cũ khi thay đổi text